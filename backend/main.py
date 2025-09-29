from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
import tempfile
from PIL import Image
import io
import subprocess
import os
import re
from docx import Document

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "Backend is working!"}

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        if not file.content_type.startswith('image/'):
            return JSONResponse(
                status_code=400,
                content={"error": "Разрешены только изображения"}
            )
        
        contents = await file.read()
        recognized_text = recognize_text_simple(contents)
        structured_data = extract_structured_data(recognized_text)
        
        return {
            "filename": file.filename,
            "content_type": file.content_type,
            "text": recognized_text,
            "characters_count": len(recognized_text),
            "structured_data": structured_data,
            "image_data": contents.hex()
        }
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Ошибка: {str(e)}"}
        )

@app.post("/export")
async def export_data(text: str = Form(...)):
    """Экспорт ТОЛЬКО в DOCX (без параметра format)"""
    try:
        doc = Document()
        doc.add_heading('Распознанный архивный документ', 0)
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=document.docx"}
        )
            
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Ошибка при создании документа: {str(e)}"}
        )

def extract_structured_data(text):
    entities = {
        'dates': [],
        'names': [],
        'archive_codes': [],
        'places': []
    }
    
    date_patterns = [
        r'\b\d{1,2}\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}\b',
        r'\b\d{4}\s+г\.',
        r'\b\d{1,2}\.\d{1,2}\.\d{4}\b'
    ]
    
    archive_patterns = [
        r'[Фф]\.\s*\d+',
        r'[Оо]п\.\s*\d+', 
        r'[Дд]\.\s*\d+',
        r'[Фф]онд\s*\d+',
        r'[Ее]д\.\s*[Хх]р\.\s*\d+'
    ]
    
    name_patterns = [
        r'\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\b'
    ]
    
    for pattern in date_patterns:
        entities['dates'].extend(re.findall(pattern, text, re.IGNORECASE))
    
    for pattern in archive_patterns:
        entities['archive_codes'].extend(re.findall(pattern, text, re.IGNORECASE))
    
    for pattern in name_patterns:
        entities['names'].extend(re.findall(pattern, text, re.IGNORECASE))
    
    return entities

def clean_recognized_text(text):
    if not text:
        return text
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line and len(line) > 2:
            cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

def recognize_text_simple(image_data: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(image_data))
        if image.mode in ['RGBA', 'P', 'LA']:
            image = image.convert('RGB')
        width, height = image.size
        if width < 1000:
            image = image.resize((width*2, height*2), Image.Resampling.LANCZOS)
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            image.save(temp_file.name, 'JPEG', quality=90)
            temp_filename = temp_file.name
        result = subprocess.run([
            'tesseract', temp_filename, 'stdout', 
            '-l', 'rus',
            '--psm', '6'
        ], capture_output=True, text=True, timeout=30)
        if os.path.exists(temp_filename):
            os.unlink(temp_filename)
        if result.returncode == 0:
            text = result.stdout.strip()
            text = clean_recognized_text(text)
            # УДАЛЕНО: блок с исправлениями
            return text if text else "Текст не распознан"
        else:
            return f"Ошибка tesseract: {result.stderr}"
    except Exception as e:
        return f"Ошибка обработки: {str(e)}"


@app.get("/health")
async def health_check():
    return {"status": "ok"}
